# import requests
# import re
# import json
# import time
# import threading
# from docx import Document
# import fitz  # PyMuPDF
# from pathlib import Path
# import os

# # ✅ Replace this with the URL printed from Colab
# COLAB_API_URL = "https://9a97-34-142-142-113.ngrok-free.app/generate"

# # === LLaMA API CALL ===
# def llama_generate(prompt, max_tokens=300, temperature=0.7):
#     try:
#         response = requests.post(COLAB_API_URL, json={
#             "prompt": prompt,
#             "max_tokens": max_tokens,
#             "temperature": temperature
#         })
#         response.raise_for_status()
#         data = response.json()
#         return data.get("response", "")
#     except Exception as e:
#         print(f"LLaMA API error: {e}")
#         return ""

# # === Input Configuration ===
# level_durations = {"easy": 5 * 60, "moderate": 10 * 60, "experienced": 15 * 60}
# level = "easy"
# interview_duration = level_durations.get(level, 10 * 60)
# level_minutes = interview_duration // 60



# # === Extract Text from PDF ===
# def extract_text_from_pdf(pdf_path):
#     doc = fitz.open(pdf_path)
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text

# # === Resume and Job Description Text ===
# current_dir = os.path.dirname(__file__)
# interview_dir = os.path.join(current_dir, "interview_doc")


# # === Extracted Text ===
# resume_text = extract_text_from_pdf(os.path.join(interview_dir, "Nikita_Pawale_8080756853.pdf"))
# job_desc_text = extract_text_from_pdf(os.path.join(interview_dir, "Data_Scientist_JD.pdf"))

# # === Generate Resume Summary using LLaMA ===
# def generate_resume_summary(resume_text):
#     prompt =  f"""Summarize this resume focusing on key skills, experience, and tools used. Write a brief, structured summary:
# {resume_text}
# Return summary:"""
#     summary= llama_generate(prompt, max_tokens=120)
#     return summary.strip().split("Return summary:")[-1].strip()
    
# resume_summary = generate_resume_summary(resume_text)


# # === Mock Data ===
# def get_level_prompt(level):
#     if level == "easy":
#         return "Ask basic, beginner-friendly interview questions."
#     elif level == "moderate":
#         return "Ask intermediate-level questions focusing on understanding and practical experience."
#     elif level == "experienced":
#         return "Ask advanced-level questions about architecture, decision-making, and deep technical topics."
#     return "Ask general technical questions."

# # === Follow-up Instruction ===
# follow_up_instruction = "Ask relevant follow-up questions based on the candidate's previous answers,job description, and resume summary. Ensure the conversation flows naturally and builds on previous responses."

# # === Base Prompt ===
# base_prompt = f"""
# You are a friendly and professional interview bot conducting a {level} level technical interview with a candidate.
# Start the conversation with a polite greeting and a brief introduction about the interview process.
# Then smoothly transition into the first technical question.
# Keep a natural, conversational tone while being clear and helpful.

# Rules:
# - Start the conversation naturaly.
# - Ask one question at a time.
# - Wait for the candidate's answer before continuing.
# - Use professional tone.
# - Total interview time: {level_minutes} minutes.
# - The conversation should feel natural and human-like.
# - Start with simple warm-up questions.
# - Gradually increase technical depth (based on {level} level).
# - As time nears the end, ask smoother conversational questions like:
#   - “Do you have any questions for us?”
#   - “Any final thoughts you'd like to share?”
#   - “Thank you for your time.”
# - {follow_up_instruction}

# ### Important:
# - Finish the interview within {level_minutes} minutes.
# - End it naturally without suddenly stopping.
# - Your role is only the Interviewer. Never answer as the candidate.

# Job Description:
# {job_desc_text}

# Candidate Resume Summary:
# {resume_summary}

# Begin the interview.

# Interviewer: Hello, thank you for joining the interview today. Let's begin. Can you briefly introduce yourself?
# """

# # === Helper: Get Answer with Timeout ===
# def get_user_input_with_reminder(prompt, timeout=60):
#     print(prompt)
#     start = time.time()

#     def user_input():
#         nonlocal answer
#         try:
#             answer = input()
#         except:
#             answer = None

#     answer = None
#     thread = threading.Thread(target=user_input)
#     thread.daemon = True
#     thread.start()
#     thread.join(timeout)

#     if answer:
#         return answer

#     print("\n⏳ Still waiting for your answer... Please respond.")
#     thread.join(timeout)

#     return answer

# # === Helper: Evaluate Answer ===
# def evaluate_answer(answer, question):
#     eval_prompt = f"""
# You are an interview evaluator. Only respond with a JSON object.

# Evaluate this answer:
# Question: {question}
# Answer: {answer}

# Return format:
# {{
#   "relevance": 0-5,
#   "technical_correctness": 0-5,
#   "clarity": 0-5,
#   "comment": "short feedback"
# }}
# """
#     raw = llama_generate(eval_prompt)
#     match = re.search(r"\{.*?\}", raw, re.DOTALL)
#     try:
#         return json.loads(match.group()) if match else {}
#     except:
#         return {"relevance": 0, "technical_correctness": 0, "clarity": 0, "comment": "Could not evaluate."}

# # === Helper: Generate Next Question ===
# def generate_next_question(history):
#     prompt = f"""
# You are a professional interviewer AI conducting a job interview.

# Context:
# Job Description: {job_desc_text}
# Candidate Resume Summary: {resume_summary}

# Rules:
# - Ask one question at a time.
# - Wait for the candidate's answer before continuing.
# - Ask questions based on the job description and the candidate's resume.
# - {get_level_prompt(level)}
# - {follow_up_instruction}

# Conversation so far:
# {history}

# Interviewer:"""
#     response = llama_generate(prompt, max_tokens=200, temperature=0.7)
#     generated = response.replace(prompt, "").strip()
#     match = re.search(r"^(.*?)(?:\n|$)", generated)
#     return match.group(1).strip() if match else generated

# # === Helper: Generate Word Report ===
# def generate_docx_report(evaluations, filename="reports/interview_report.docx"):
#     doc = Document()
#     doc.add_heading("Interview Report", 0)

#     for i, e in enumerate(evaluations):
#         doc.add_heading(f"Q{i+1}: {e['question']}", level=2)
#         doc.add_paragraph(f"Answer: {e['answer']}")
#         doc.add_paragraph(
#             f"Relevance: {e['relevance']}, Technical: {e['technical_correctness']}, Clarity: {e['clarity']}"
#         )
#         doc.add_paragraph(f"Comment: {e['comment']}\n")

#     avg_rel = sum(int(e['relevance']) for e in evaluations) / len(evaluations) if evaluations else 0
#     avg_tech = sum(int(e['technical_correctness']) for e in evaluations) / len(evaluations) if evaluations else 0
#     avg_clarity = sum(int(e['clarity']) for e in evaluations) / len(evaluations) if evaluations else 0

#     doc.add_heading("Average Scores", level=2)
#     doc.add_paragraph(
#         f"Relevance: {avg_rel:.2f}, Technical: {avg_tech:.2f}, Clarity: {avg_clarity:.2f}"
#     )

#     doc.save(filename)
#     print(f"\n📝 Interview report saved as: {filename}")

# # === Extract Text from PDF ===
# def extract_text_from_pdf(pdf_path):
#     doc = fitz.open(pdf_path)
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text



# # === Interview Start ===
# def run_interview():
#     if input("\nStart the interview? (y/n): ").lower() != "y":
#         print("Interview canceled.")
#         return

#     interview_start = time.time()
#     interview_end = interview_start + interview_duration
#     conversation = base_prompt
#     qa_log = []  # Store (question, answer) pairs

#     # First Question
#     question = generate_next_question(conversation)
#     print("\n🧑‍💼 Interviewer:", question)
#     conversation += f"\nInterviewer: {question}"

#     # Main Interview Loop
#     WRAP_UP_THRESHOLD = 120
#     in_wrap_up = False

#     while time.time() < interview_end:
#         remaining = interview_end - time.time()
#         if remaining < WRAP_UP_THRESHOLD and not in_wrap_up:
#             print("\n⏳ Time is almost up. Wrapping up soon...")
#             conversation += "\n[NOTE: Wrapping up the interview.]"
#             in_wrap_up = True

#         answer = get_user_input_with_reminder("\n🧑 Candidate (your answer): ")
#         if not answer:
#             print("\n🛑 No answer received. Ending interview.")
#             break

#         conversation += f"\nCandidate: {answer}"
#         qa_log.append({"question": question, "answer": answer})  # Store Q&A

#         question = generate_next_question(conversation)
#         if not question:
#             print("\n🛑 No further questions generated. Ending interview.")
#             break

#         print("\n🧑‍💼 Interviewer:", question)
#         conversation += f"\nInterviewer: {question}"

#     print("\n✅ Interview complete.")

#     # Evaluate all answers after the interview
#     evaluation_log = []
#     for qa in qa_log:
#         eval_result = evaluate_answer(qa["answer"], qa["question"])
#         eval_result["question"] = qa["question"]
#         eval_result["answer"] = qa["answer"]
#         evaluation_log.append(eval_result)

#     generate_docx_report(evaluation_log)

# # To run the interview, just call:
# # run_interview()
import requests
import re
import json
import time
import threading
from docx import Document
import fitz  # PyMuPDF
from pathlib import Path
import os

# ✅ Replace with your Mistral API key
MISTRAL_API_KEY = "KZEcWkcI44L0T7t734LVe1FfImLG9VnO"
MISTRAL_API_URL = "https://api.mistral.ai/v1/chat/completions"

# === Mistral API CALL ===
def mistral_generate(prompt, max_tokens=300, temperature=0.7, model="mistral-small-latest"):
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MISTRAL_API_KEY}"
    }
    
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You are a professional interview bot conducting technical interviews."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": temperature
    }
    
    try:
        response = requests.post(MISTRAL_API_URL, headers=headers, json=payload)
        response.raise_for_status()  # Raises an HTTPError for bad responses
        
        data = response.json()
        return data["choices"][0]["message"]["content"].strip()
    except requests.exceptions.RequestException as e:
        print(f"Mistral API request error: {e}")
        return ""
    except KeyError as e:
        print(f"Mistral API response format error: {e}")
        return ""
    except Exception as e:
        print(f"Mistral API error: {e}")
        return ""

# === Input Configuration ===
level_durations = {"easy": 5 * 60, "moderate": 10 * 60, "experienced": 15 * 60}
level = "easy"
interview_duration = level_durations.get(level, 10 * 60)
level_minutes = interview_duration // 60

# === Extract Text from PDF ===
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# === Resume and Job Description Text ===
current_dir = os.path.dirname(__file__)
interview_dir = os.path.join(current_dir, "interview_doc")

# === Extracted Text ===
# Fix the file paths - remove the absolute path duplicates
resume_path = os.path.join(interview_dir, "Resume (4).pdf")
job_desc_path = os.path.join(interview_dir, "JD.txt")

# Check if files exist before processing
if os.path.exists(resume_path):
    resume_text = extract_text_from_pdf(resume_path)
else:
    print(f"Resume file not found: {resume_path}")
    resume_text = "Resume not available"

if os.path.exists(job_desc_path):
    with open(job_desc_path, 'r', encoding='utf-8') as f:
        job_desc_text = f.read()
else:
    print(f"Job description file not found: {job_desc_path}")
    job_desc_text = "Job description not available"

# === Generate Resume Summary using Mistral ===
def generate_resume_summary(resume_text):
    prompt = f"""Summarize this resume focusing on key skills, experience, and tools used. Write a brief, structured summary:
{resume_text}
Return summary:"""
    summary = mistral_generate(prompt, max_tokens=120)
    return summary.strip().split("Return summary:")[-1].strip()
    
resume_summary = generate_resume_summary(resume_text)

# === Mock Data ===
def get_level_prompt(level):
    if level == "easy":
        return "Ask basic, beginner-friendly interview questions."
    elif level == "moderate":
        return "Ask intermediate-level questions focusing on understanding and practical experience."
    elif level == "experienced":
        return "Ask advanced-level questions about architecture, decision-making, and deep technical topics."
    return "Ask general technical questions."

# === Follow-up Instruction ===
follow_up_instruction = "Ask relevant follow-up questions based on the candidate's previous answers, job description, and resume summary. Ensure the conversation flows naturally and builds on previous responses."

# === Base Prompt ===
base_prompt = f"""
You are a friendly and professional interview bot conducting a {level} level technical interview with a candidate.
Start the conversation with a polite greeting and a brief introduction about the interview process.
Then smoothly transition into the first technical question.
Keep a natural, conversational tone while being clear and helpful.

Rules:
- Start the conversation naturally.
- Ask one question at a time.
- Wait for the candidate's answer before continuing.
- Use professional tone.
- Total interview time: {level_minutes} minutes.
- The conversation should feel natural and human-like.
- Start with simple warm-up questions.
- Gradually increase technical depth (based on {level} level).
- As time nears the end, ask smoother conversational questions like:
  - "Do you have any questions for us?"
  - "Any final thoughts you'd like to share?"
  - "Thank you for your time."
- {follow_up_instruction}

### Important:
- Finish the interview within {level_minutes} minutes.
- End it naturally without suddenly stopping.
- Your role is only the Interviewer. Never answer as the candidate.

Job Description:
{job_desc_text}

Candidate Resume Summary:
{resume_summary}

Begin the interview.

Interviewer: Hello, thank you for joining the interview today. Let's begin. Can you briefly introduce yourself?
"""

# === Helper: Get Answer with Timeout ===
def get_user_input_with_reminder(prompt, timeout=60):
    print(prompt)
    start = time.time()

    def user_input():
        nonlocal answer
        try:
            answer = input()
        except:
            answer = None

    answer = None
    thread = threading.Thread(target=user_input)
    thread.daemon = True
    thread.start()
    thread.join(timeout)

    if answer:
        return answer

    print("\n⏳ Still waiting for your answer... Please respond.")
    thread.join(timeout)

    return answer

# === Helper: Evaluate Answer ===
def evaluate_answer(answer, question):
    eval_prompt = f"""
You are an interview evaluator. Only respond with a JSON object.

Evaluate this answer:
Question: {question}
Answer: {answer}

Return format:
{{
  "relevance": 0-5,
  "technical_correctness": 0-5,
  "clarity": 0-5,
  "comment": "short feedback"
}}
"""
    raw = mistral_generate(eval_prompt, max_tokens=150)
    match = re.search(r"\{.*?\}", raw, re.DOTALL)
    try:
        return json.loads(match.group()) if match else {}
    except:
        return {"relevance": 0, "technical_correctness": 0, "clarity": 0, "comment": "Could not evaluate."}

# === Helper: Generate Next Question ===
def generate_next_question(history):
    prompt = f"""
You are a professional interviewer AI conducting a job interview.

Context:
Job Description: {job_desc_text}
Candidate Resume Summary: {resume_summary}

Rules:
- Ask one question at a time.
- Wait for the candidate's answer before continuing.
- Ask questions based on the job description and the candidate's resume.
- {get_level_prompt(level)}
- {follow_up_instruction}

Conversation so far:
{history}

Interviewer:"""
    response = mistral_generate(prompt, max_tokens=200, temperature=0.7)
    generated = response.replace(prompt, "").strip()
    match = re.search(r"^(.*?)(?:\n|$)", generated)
    return match.group(1).strip() if match else generated

# === Helper: Generate Word Report ===
def generate_docx_report(evaluations, filename="reports/interview_report.docx"):
    # Create reports directory if it doesn't exist
    os.makedirs("reports", exist_ok=True)
    
    doc = Document()
    doc.add_heading("Interview Report", 0)

    for i, e in enumerate(evaluations):
        doc.add_heading(f"Q{i+1}: {e['question']}", level=2)
        doc.add_paragraph(f"Answer: {e['answer']}")
        doc.add_paragraph(
            f"Relevance: {e['relevance']}, Technical: {e['technical_correctness']}, Clarity: {e['clarity']}"
        )
        doc.add_paragraph(f"Comment: {e['comment']}\n")

    if evaluations:
        avg_rel = sum(int(e.get('relevance', 0)) for e in evaluations) / len(evaluations)
        avg_tech = sum(int(e.get('technical_correctness', 0)) for e in evaluations) / len(evaluations)
        avg_clarity = sum(int(e.get('clarity', 0)) for e in evaluations) / len(evaluations)

        doc.add_heading("Average Scores", level=2)
        doc.add_paragraph(
            f"Relevance: {avg_rel:.2f}, Technical: {avg_tech:.2f}, Clarity: {avg_clarity:.2f}"
        )

    doc.save(filename)
    print(f"\n📝 Interview report saved as: {filename}")

# === Interview Start ===
def run_interview():
    if input("\nStart the interview? (y/n): ").lower() != "y":
        print("Interview canceled.")
        return

    interview_start = time.time()
    interview_end = interview_start + interview_duration
    conversation = base_prompt
    qa_log = []  # Store (question, answer) pairs

    # First Question
    question = generate_next_question(conversation)
    print("\n🧑‍💼 Interviewer:", question)
    conversation += f"\nInterviewer: {question}"

    # Main Interview Loop
    WRAP_UP_THRESHOLD = 120
    in_wrap_up = False

    while time.time() < interview_end:
        remaining = interview_end - time.time()
        if remaining < WRAP_UP_THRESHOLD and not in_wrap_up:
            print("\n⏳ Time is almost up. Wrapping up soon...")
            conversation += "\n[NOTE: Wrapping up the interview.]"
            in_wrap_up = True

        answer = get_user_input_with_reminder("\n🧑 Candidate (your answer): ")
        if not answer:
            print("\n🛑 No answer received. Ending interview.")
            break

        conversation += f"\nCandidate: {answer}"
        qa_log.append({"question": question, "answer": answer})  # Store Q&A

        question = generate_next_question(conversation)
        if not question:
            print("\n🛑 No further questions generated. Ending interview.")
            break

        print("\n🧑‍💼 Interviewer:", question)
        conversation += f"\nInterviewer: {question}"

    print("\n✅ Interview complete.")

    # Evaluate all answers after the interview
    evaluation_log = []
    for qa in qa_log:
        eval_result = evaluate_answer(qa["answer"], qa["question"])
        eval_result["question"] = qa["question"]
        eval_result["answer"] = qa["answer"]
        evaluation_log.append(eval_result)

    generate_docx_report(evaluation_log)

# === Main execution ===
if __name__ == "__main__":
    print("🤖 Interview Bot (Mistral API Version)")
    print(f"📊 Level: {level}")
    print(f"⏱️  Duration: {level_minutes} minutes")
    print(f"📄 Resume Summary: {resume_summary[:100]}...")
    run_interview()